from docx import Document
from docx import *
import os
#print("Enter number of files:")
#nof=input()
#nof=int(nof)
file_names=[]
nof=0
#print("Enter file names one by one!")
#for i in range(0,nof):
#	file_names.append(input())
file=open("docx_files.txt","r")
file_names=file.readlines()
nof=len(file_names)
#print(file_names[0][:-6])
for i in range(0,nof):
	file_name=file_names[i][:-6]
	document=Document(file_name+".docx")
	name='Raja Vijay Parmar'
	roll_no='14102A0005'
	tables=document.tables
	have_to_change_name=0
	have_to_change_rn=0
	for table in tables:
		for row in table.rows:
			for cell in row.cells:
				for p in cell.paragraphs:
					if(have_to_change_name==1):
						p.text=name
						have_to_change_name=0
					if(have_to_change_rn==1):
						have_to_change_rn=0
						p.text=roll_no
					if(p.text=='Student Name'):
						have_to_change_name=1
					if(p.text=='Roll Number'):
						have_to_change_rn=1
	document.save(file_name+"_r_.docx")
	os.system("mv *_r_*.docx /srv/ftp")
	